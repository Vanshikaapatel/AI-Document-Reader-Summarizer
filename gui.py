# gui.py

import sys
import os
import threading
import traceback

# Prepare import paths
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
if BASE_DIR not in sys.path:
    sys.path.insert(0, BASE_DIR)
SUMM_DIR = os.path.join(BASE_DIR, "summarizer")
if SUMM_DIR not in sys.path:
    sys.path.insert(0, SUMM_DIR)

# Tkinter & Helpers
import tkinter as tk
from tkinter import ttk, filedialog, messagebox, scrolledtext

# Summarizer modules
from summarizer.pdf_reader import extract_text_from_pdf
from summarizer.text_preprocess import clean_text
from summarizer.summarize_model import summarize_text

# Optional TTS
try:
    import pyttsx3
    TTS_AVAILABLE = True
except:
    TTS_AVAILABLE = False

# PDF Export
try:
    from reportlab.pdfgen import canvas
    PDF_AVAILABLE = True
except:
    PDF_AVAILABLE = False

# DOCX Export
try:
    from docx import Document
    DOCX_AVAILABLE = True
except:
    DOCX_AVAILABLE = False


class SummarizerApp:
    def __init__(self, root):
        self.root = root
        root.title("AI Document Summarizer")
        root.geometry("1050x720")

        self.is_dark_mode = False  # default light mode
        self.spinner_running = False
        self.spinner_frames = ["⠋", "⠙", "⠹", "⠸", "⠼", "⠴", "⠦", "⠧", "⠇", "⠏"]

        # ------------ TOP BAR ------------
        top = ttk.Frame(root)
        top.pack(fill=tk.X, padx=10, pady=10)

        self.file_var = tk.StringVar()
        ttk.Entry(top, textvariable=self.file_var, width=80).pack(side=tk.LEFT, padx=4)
        ttk.Button(top, text="Browse PDF", command=self.browse_file).pack(side=tk.LEFT, padx=4)

        # ----------- Dark Mode Toggle -----------
        ttk.Button(top, text="🌙 Dark Mode", command=self.toggle_dark_mode).pack(side=tk.LEFT, padx=10)

        # ------------ OPTIONS ----------------
        opts = ttk.Frame(root)
        opts.pack(fill=tk.X, padx=10)

        self.mode = tk.StringVar(value="abstractive")
        ttk.Radiobutton(opts, text="Abstractive", value="abstractive", variable=self.mode).pack(side=tk.LEFT)
        ttk.Radiobutton(opts, text="Extractive", value="extractive", variable=self.mode).pack(side=tk.LEFT)

        self.ocr = tk.BooleanVar()
        ttk.Checkbutton(opts, text="Use OCR", variable=self.ocr).pack(side=tk.LEFT, padx=10)

        ttk.Label(opts, text="Summary length:").pack(side=tk.LEFT, padx=(20, 5))
        self.len_var = tk.IntVar(value=200)
        ttk.Spinbox(opts, from_=50, to=2000, increment=50, textvariable=self.len_var, width=8).pack(side=tk.LEFT)

        # ------------ BUTTON BAR ------------
        btns = ttk.Frame(root)
        btns.pack(fill=tk.X, padx=10, pady=8)

        ttk.Button(btns, text="Generate Summary", command=self.run_summary).pack(side=tk.LEFT, padx=4)
        ttk.Button(btns, text="Save as TXT", command=self.export_txt).pack(side=tk.LEFT, padx=4)
        ttk.Button(btns, text="Save as PDF", command=self.export_pdf).pack(side=tk.LEFT, padx=4)
        ttk.Button(btns, text="Save as DOCX", command=self.export_docx).pack(side=tk.LEFT, padx=4)

        self.speak_btn = ttk.Button(btns, text="Speak Summary", command=self.speak_summary)
        self.speak_btn.pack(side=tk.LEFT, padx=4)
        if not TTS_AVAILABLE:
            self.speak_btn.config(state="disabled", text="Install pyttsx3")

        # -------------- LOADING SPINNER --------------
        self.status_var = tk.StringVar(value="")
        ttk.Label(root, textvariable=self.status_var).pack()
        self.spinner_label = ttk.Label(root, text="", font=("Segoe UI", 16))
        self.spinner_label.pack()

        # -------------- NOTEBOOK (tabs) --------------
        self.notebook = ttk.Notebook(root)
        self.notebook.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

        # Raw Text Tab
        self.raw_text = scrolledtext.ScrolledText(self.notebook, wrap=tk.WORD, font=("Segoe UI", 11))
        self.notebook.add(self.raw_text, text="Extracted Text")

        # Summary Tab
        self.summary_box = scrolledtext.ScrolledText(self.notebook, wrap=tk.WORD, font=("Segoe UI", 11))
        self.notebook.add(self.summary_box, text="Summary")

        self.summary_content = ""
        self.raw_content = ""

    # ----------------- DARK MODE FUNCTION -----------------
    def toggle_dark_mode(self):
        self.is_dark_mode = not self.is_dark_mode

        bg = "#1e1e1e" if self.is_dark_mode else "white"
        fg = "white" if self.is_dark_mode else "black"

        self.root.configure(bg=bg)
        self.raw_text.configure(bg=bg, fg=fg, insertbackground=fg)
        self.summary_box.configure(bg=bg, fg=fg, insertbackground=fg)
        self.spinner_label.configure(background=bg, foreground=fg)

    # ----------------- FILE BROWSER -----------------
    def browse_file(self):
        path = filedialog.askopenfilename(filetypes=[("PDF Files", "*.pdf")])
        if path:
            self.file_var.set(path)

    # ----------------- MAIN RUNNER -----------------
    def run_summary(self):
        pdf = self.file_var.get()
        if not pdf:
            messagebox.showwarning("No File", "Select a PDF.")
            return

        self.start_spinner("Extracting and summarizing...")
        threading.Thread(target=self.worker, args=(pdf,), daemon=True).start()

    # ----------------- WORKER THREAD -----------------
    def worker(self, pdf):
        try:
            raw = extract_text_from_pdf(pdf, use_ocr=self.ocr.get())
            self.raw_content = raw
            self.safe_update_raw(raw)

            cleaned = clean_text(raw)
            result = summarize_text(cleaned, mode=self.mode.get(), max_length=self.len_var.get())

            self.summary_content = result
            self.safe_update_summary(result)

        except Exception as e:
            tb = traceback.format_exc()
            self.safe_update_summary(f"Error:\n{e}\n\n{tb}")

        self.stop_spinner("Done ✔")

    # ----------------- SPINNER -----------------
    def start_spinner(self, message):
        self.spinner_running = True
        self.status_var.set(message)
        self.animate_spinner()

    def animate_spinner(self):
        if not self.spinner_running:
            self.spinner_label.config(text="")
            return
        frame = self.spinner_frames.pop(0)
        self.spinner_frames.append(frame)
        self.spinner_label.config(text=frame)
        self.root.after(120, self.animate_spinner)

    def stop_spinner(self, message):
        self.spinner_running = False
        self.status_var.set(message)

    # ----------------- SAFE UI UPDATE -----------------
    def safe_update_raw(self, text):
        self.raw_text.after(0, lambda: self._update_text_widget(self.raw_text, text))

    def safe_update_summary(self, text):
        self.summary_box.after(0, lambda: self._update_text_widget(self.summary_box, text))

    @staticmethod
    def _update_text_widget(widget, text):
        widget.delete("1.0", tk.END)
        widget.insert(tk.END, text)

    # ----------------- EXPORT FEATURES -----------------
    def export_txt(self):
        if not self.summary_content:
            messagebox.showinfo("Empty", "Generate summary first.")
            return

        file = filedialog.asksaveasfilename(defaultextension=".txt")
        if file:
            with open(file, "w", encoding="utf-8") as f:
                f.write(self.summary_content)
            messagebox.showinfo("Saved", "Saved as TXT")

    def export_pdf(self):
        if not PDF_AVAILABLE:
            messagebox.showerror("Missing Package", "Install reportlab:\npip install reportlab")
            return

        file = filedialog.asksaveasfilename(defaultextension=".pdf")
        if file:
            c = canvas.Canvas(file)
            text = c.beginText(40, 800)
            for line in self.summary_content.split("\n"):
                text.textLine(line)
            c.drawText(text)
            c.save()
            messagebox.showinfo("Saved", "Saved as PDF")

    def export_docx(self):
        if not DOCX_AVAILABLE:
            messagebox.showerror("Missing Package", "Install python-docx:\npip install python-docx")
            return

        file = filedialog.asksaveasfilename(defaultextension=".docx")
        if file:
            doc = Document()
            doc.add_heading("Summary", level=1)
            doc.add_paragraph(self.summary_content)
            doc.save(file)
            messagebox.showinfo("Saved", "Saved as DOCX")

    # ----------------- TEXT TO SPEAK -----------------
    def speak_summary(self):
        if not self.summary_content:
            messagebox.showinfo("Empty", "Generate a summary first.")
            return
        if not TTS_AVAILABLE:
            messagebox.showerror("Missing Package", "Install pyttsx3")
            return

        threading.Thread(target=self._speak_worker, daemon=True).start()

    def _speak_worker(self):
        engine = pyttsx3.init()
        engine.say(self.summary_content)
        engine.runAndWait()


def main():
    root = tk.Tk()
    SummarizerApp(root)
    root.mainloop()


if __name__ == "__main__":
    main()
