import streamlit as st
import sys, os, tempfile
from io import BytesIO

import streamlit as st

# Redirect if not logged in
if "user" not in st.session_state or st.session_state.user is None:
    st.warning("Please login to access the summarizer.")
    st.stop()


st.write(f"👤 Logged in as: **{st.session_state.user}**")

if st.button("🚪 Logout"):
    st.session_state.user = None
    st.success("Logged out successfully!")
    st.switch_page("App.py")

if "user" not in st.session_state:
    st.session_state.user = None


# Fix imports for summarizer package
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
sys.path.insert(0, BASE_DIR)

from summarizer.pdf_reader import extract_text_from_pdf
from summarizer.text_preprocess import clean_text
from summarizer.summarize_model import summarize_text

# Optional modules
try:
    import pyttsx3
    TTS_AVAILABLE = True
except:
    TTS_AVAILABLE = False

try:
    from reportlab.pdfgen import canvas
    PDF_AVAILABLE = True
except:
    PDF_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except:
    DOCX_AVAILABLE = False

# ---------------------------------------------------------
# PAGE CONFIG
# ---------------------------------------------------------
st.set_page_config(
    page_title="AI Document Summarizer",
    layout="wide",
    page_icon="📝"
)

# ---- HIDE ONLY SIDEBAR TOGGLE ARROW (NOT SIDEBAR) ----
st.markdown("""
    <style>

    /* Hide ONLY the sidebar collapse/expand arrow */
    button[data-testid="collapsedControl"] {
        display: none !important;
    }

    </style>
    """, unsafe_allow_html=True)



# ---- HIDE STREAMLIT SIDEBAR TOGGLE ARROW ----
st.markdown("""
    <style>
    /* Hide the sidebar toggle arrow (keyboard_double_arrow_right) */
    button[kind="header"] {
        display: none !important;
    }

    /* Extra safety for material icon text */
    .material-icons {
        display: none !important;
    }
    </style>
    """, unsafe_allow_html=True)


# ------------------- NAVBAR -------------------

def navbar(active="Summarizer"):
    st.markdown("""
    <style>
    .nav-container {
        display: flex;
        justify-content: center;
        gap: 35px;
        padding: 18px;
        margin-bottom: 15px;
    }

    .nav-btn {
        padding: 8px 22px;
        border-radius: 8px;
        background: rgba(255,255,255,0.08);
        cursor: pointer;
        color: #e0e0e0;
        font-weight: 600;
        text-decoration: none;
        transition: 0.3s;
    }
    .nav-btn:hover {
        background: rgba(255,255,255,0.18);
        transform: translateY(-2px);
    }

    .active {
        background: linear-gradient(90deg,#6a11cb,#2575fc);
        color: white !important;
        box-shadow: 0 4px 12px rgba(37,117,252,0.4);
    }
    </style>
    """, unsafe_allow_html=True)

    # ---------- NAVBAR BUTTONS ----------
    col1, col2, col3, col4, col5 = st.columns(5)

    with col1:
        if st.button("Home", key="home_btn"):
            st.switch_page("App.py")

    with col2:
        if st.button("Login", key="Login_btn"):
            st.switch_page("pages/1_Login.py")

    with col3:
        if st.button("Register", key="Register_btn"):
            st.switch_page("pages/2_Register.py")

    with col4:
        if st.button("Summarizer", key="Summarizer_btn"):
            st.switch_page("pages/3_Summarizer.py")

    with col5:
        if st.button("Admin", key="admin_btn"):
            st.switch_page("pages/4_Admin.py")
navbar("summarizer")



# ---------------------------------------------------------
# ADVANCED PROFESSIONAL WEBSITE UI STYLING
# ---------------------------------------------------------
st.markdown(
    """
    <style>

    /* Import Modern Website Font */
    @import url('https://fonts.googleapis.com/css2?family=Poppins:wght@300;400;500;600;700&display=swap');

    * {
        font-family: 'Poppins', sans-serif !important;
    }

    

    .main-container:hover {
        transform: scale(1.01);
    }

    /* Modernized Heading */
    h1 {
        font-size: 48px !important;
        font-weight: 700 !important;
        text-align: center;
        margin-bottom: 15px;
        letter-spacing: -1px;
    }

    /* Subtitle */
    .subtext {
        text-align: center;
        font-size: 20px;
        opacity: 0.85;
        margin-bottom: 25px;
    }

    /* Gradient Buttons (Professional UI) */
    .stButton>button {
        background: linear-gradient(135deg, #6a11cb, #2575fc);
        border: none;
        padding: 12px 26px;
        border-radius: 10px;
        color: white;
        font-size: 18px;
        font-weight: 500;
        transition: 0.25s;
        box-shadow: 0 4px 14px rgba(0, 102, 255, 0.35);
    }

    .stButton>button:hover {
        transform: translateY(-2px);
        box-shadow: 0 6px 20px rgba(0, 102, 255, 0.45);
        background: linear-gradient(135deg, #6a11cb, #2575fc);
    }

    /* File Upload Area Custom Design */
    .uploadedFile {
        background-color: rgba(255,255,255,0.30) !important;
        border-radius: 12px !important;
        padding: 20px !important;
        border: 2px dashed #6a11cb !important;
        transition: 0.3s;
    }

    .uploadedFile:hover {
        border-color: #2575fc !important;
    }

    /* Download Buttons */
    .stDownloadButton>button {
       background: linear-gradient(135deg, #6a11cb, #2575fc);
        padding: 10px 22px;
        border-radius: 10px;
        color: white !important;
        border: none;
        transition: 0.25s;
    }

    .stDownloadButton>button:hover {
        transform: scale(1.05);
        background: linear-gradient(135deg, #6a11cb, #2575fc);
    }

    /* Expander styling */
    .streamlit-expanderHeader {
        font-size: 18px !important;
        font-weight: 600 !important;
    }

    </style>
    """,
    unsafe_allow_html=True
)



# ---------------------------------------------------------
# THEME TOGGLE (ONLY MAIN PAGE CHANGES)
# ---------------------------------------------------------
theme = st.sidebar.radio("Theme", ["🌙 Dark Mode", "☀️ Light Mode"])

if theme == "🌙 Dark Mode":
    st.markdown("""
        <style>
        .stApp { background-color: #0d1117 !important; color: white !important; }
        .main-container { background: rgba(255,255,255,0.07) !important; color: white !important; }
        p, label, span, div, h1, h2, h3, h4 { color: white !important; }
        </style>
    """, unsafe_allow_html=True)

else:
    st.markdown("""
        <style>
        .stApp { background-color: #f4f4f4 !important; color: black !important; }
        .main-container { background: rgba(255,255,255,0.07) !important; color: black !important; }
        p, label, span, div, h1, h2, h3, h4 { color: black !important; }
        </style>
    """, unsafe_allow_html=True)



# ---------------------------------------------------------
# HEADER
# ---------------------------------------------------------
st.markdown("<h1>📝 AI Document Reader & Summarizer</h1>", unsafe_allow_html=True)
st.markdown("<div class='subtext'>Upload → Extract → Summarize → Export with AI</div>", unsafe_allow_html=True)



# ---------------------------------------------------------
# SIDEBAR OPTIONS
# ---------------------------------------------------------
st.sidebar.subheader("Summary Options")

mode = st.sidebar.selectbox("Summarization Mode", ["abstractive", "extractive"])
enable_ocr = st.sidebar.checkbox("Enable OCR for scanned PDFs")
max_len = st.sidebar.slider("Summary Length (words)", 50, 500, 150, step=25)

if st.session_state.user == "admin":
    if st.sidebar.button("🔐 Admin"):
        st.switch_page("pages/4_Admin.py")




# ---------------------------------------------------------
# MAIN CONTENT (NOW STYLED)
# ---------------------------------------------------------
with st.container():
    st.markdown("<div class='main-container'>", unsafe_allow_html=True)

    uploaded_file = st.file_uploader("📂 Upload PDF File", type=["pdf"], help="Upload your document here")

    if uploaded_file:
        st.info("📄 Processing your PDF...")

        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
            tmp.write(uploaded_file.read())
            pdf_path = tmp.name

        # Extract text
        with st.spinner("🔍 Extracting text..."):
            raw_text = extract_text_from_pdf(pdf_path, use_ocr=enable_ocr)
            from summarizer.document_classifier import detect_document_type
            doc_type = detect_document_type(raw_text)
            st.info(f"📄 Detected Document Type: **{doc_type}**")

        if not raw_text.strip():
            st.error("⚠ No readable text found. Try enabling OCR.")
        else:
            with st.expander("📄 Extracted Text (Preview)"):
                st.text(raw_text[:6000])

            cleaned = clean_text(raw_text)

            # Summarization
            with st.spinner("🤖 Generating Summary..."):
                summary = summarize_text(cleaned, mode=mode, max_length=max_len)

            st.subheader("📌 Summary Output")
            st.write(summary)


            # -----------------------------------------------------
            # TEXT-TO-SPEECH AUDIO
            # -----------------------------------------------------
            if TTS_AVAILABLE:
                st.write("### 🔊 Listen to Summary")

                engine = pyttsx3.init()
                engine.save_to_file(summary, "summary_audio.mp3")
                engine.runAndWait()

                audio_file = open("summary_audio.mp3", "rb").read()

                st.audio(audio_file, format="audio/mp3")

                st.download_button(
                    "⬇️ Download Audio",
                    audio_file,
                    file_name="summary_audio.mp3"
                )


            # -----------------------------------------------------
            # EXPORT OPTIONS
            # -----------------------------------------------------
            st.write("### 💾 Download Summary")

            # TXT
            st.download_button("📄 Download TXT", summary, file_name="summary.txt")

            # PDF
            if PDF_AVAILABLE:
                pdf_buffer = BytesIO()
                c = canvas.Canvas(pdf_buffer)
                text_obj = c.beginText(40, 800)
                for line in summary.split("\n"):
                    text_obj.textLine(line)
                c.drawText(text_obj)
                c.save()
                
                # Increase user upload count
                users = json.load(open("users.json"))
                users[st.session_state.user]["upload_count"] += 1
                json.dump(users, open("users.json", "w"), indent=4)

                st.download_button(
                    "📘 Download PDF",
                    data=pdf_buffer.getvalue(),
                    file_name="summary.pdf",
                    mime="application/pdf"
                )
                
               # Increase summary count
                users = json.load(open("users.json"))
                users[st.session_state.user]["summary_count"] += 1
                json.dump(users, open("users.json", "w"), indent=4)


            # DOCX
            if DOCX_AVAILABLE:
                doc = Document()
                doc.add_heading("AI Summary", level=1)
                doc.add_paragraph(summary)

                doc_buffer = BytesIO()
                doc.save(doc_buffer)

                st.download_button(
                    "📄 Download DOCX",
                    data=doc_buffer.getvalue(),
                    file_name="summary.docx"
                )

        os.remove(pdf_path)

    st.markdown("</div>", unsafe_allow_html=True)
